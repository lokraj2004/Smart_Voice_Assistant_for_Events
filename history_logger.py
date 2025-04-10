import threading
from docx import Document
from datetime import datetime

class HistoryLogger:
    def __init__(self, file_path="interaction_history.docx"):
        self.file_path = file_path
        self.document = Document()
        self.document.add_heading('User Interaction History', 0)
        self.lock = threading.Lock()
        self.logging_active = False

    def start_logging(self):
        with self.lock:
            if not self.logging_active:
                self.logging_active = True
                self._log_message("🔴 Logging started")

    def stop_logging(self):
        with self.lock:
            if self.logging_active:
                self._log_message("🟢 Logging stopped")
                self.logging_active = False
                self.save()

    def log(self, user_input, response=None):
        if not self.logging_active:
            return
        with self.lock:
            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            self.document.add_paragraph(f"[{timestamp}] 🧍 You: {user_input}")
            if response:
                self.document.add_paragraph(f"[{timestamp}] 🤖 Assistant: {response}")

    def save(self):
        with self.lock:
            self.document.save(self.file_path)

    def _log_message(self, msg):
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        self.document.add_paragraph(f"[{timestamp}] {msg}")
